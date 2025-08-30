# -*- coding: utf-8 -*-
import os
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors

# === กำหนดโฟลเดอร์ปลายทาง ===
OUT_DIR = r"C:\ThaiLesson"   # เปลี่ยนได้ตามต้องการ
os.makedirs(OUT_DIR, exist_ok=True)

# -------------------------------
# ส่วนที่ 1 : สร้าง DOCX
# -------------------------------
docx_path = os.path.join(OUT_DIR, "แบบเรียน_ฝึกอ่าน_สี_วัน_เดือน.docx")
doc = Document()

# ฟอนต์ไทย
style = doc.styles['Normal']
style.font.name = 'TH Sarabun New'
style._element.rPr.rFonts.set(qn('w:eastAsia'), 'TH Sarabun New')
style.font.size = Pt(16)

doc.add_heading("แบบเรียนฝึกอ่านภาษาไทย: เรื่องสี – วัน – เดือน", 0)

# Unit สี
doc.add_heading("Unit 1: สี", level=1)
doc.add_paragraph("คำศัพท์: แดง, เหลือง, น้ำเงิน, เขียว, ส้ม, ชมพู, ม่วง, น้ำตาล, ขาว, ดำ, เทา, ฟ้า")
doc.add_paragraph("ประโยคตัวอย่าง: นี่คือ…สี…, …ของฉันสี…, ฉันชอบสี…")

# Unit วัน
doc.add_heading("Unit 2: วัน", level=1)
doc.add_paragraph("วันทั้ง 7: อาทิตย์, จันทร์, อังคาร, พุธ, พฤหัสบดี, ศุกร์, เสาร์")
doc.add_paragraph("ประโยคตัวอย่าง: วันนี้วัน…, พรุ่งนี้วัน…, เมื่อวานวัน…")

# Unit เดือน
doc.add_heading("Unit 3: เดือน", level=1)
doc.add_paragraph("12 เดือน: มกราคม–ธันวาคม พร้อมตัวย่อ ม.ค.–ธ.ค.")
doc.add_paragraph("ประโยคตัวอย่าง: เดือนนี้คือเดือน…, ฉันเกิดเดือน…")

# ปฏิทินเปล่า
doc.add_heading("ใบงาน: ปฏิทินเปล่า", level=1)
doc.add_paragraph("เดือน………………… ปี………………… เขียนตัวเลขวันที่ลงในช่องตาราง")

doc.save(docx_path)

# -------------------------------
# ส่วนที่ 2 : สร้าง PDF
# -------------------------------
pdf_path = os.path.join(OUT_DIR, "แบบเรียน_ฝึกอ่าน_สี_วัน_เดือน.pdf")

styles = getSampleStyleSheet()
styles.add(ParagraphStyle(name='TitleTH', fontName='Helvetica-Bold', fontSize=20, alignment=1))
styles.add(ParagraphStyle(name='HeadingTH', fontName='Helvetica-Bold', fontSize=14))
styles.add(ParagraphStyle(name='NormalTH', fontName='Helvetica', fontSize=12))

story = []
story.append(Paragraph("แบบเรียนฝึกอ่านภาษาไทย: เรื่องสี – วัน – เดือน", styles['TitleTH']))
story.append(Spacer(1, 20))

story.append(Paragraph("Unit 1: สี", styles['HeadingTH']))
story.append(Paragraph("แดง, เหลือง, น้ำเงิน, เขียว, ส้ม, ชมพู, ม่วง, น้ำตาล, ขาว, ดำ, เทา, ฟ้า", styles['NormalTH']))
story.append(PageBreak())

story.append(Paragraph("Unit 2: วัน", styles['HeadingTH']))
story.append(Paragraph("อาทิตย์, จันทร์, อังคาร, พุธ, พฤหัสบดี, ศุกร์, เสาร์", styles['NormalTH']))
story.append(PageBreak())

story.append(Paragraph("Unit 3: เดือน", styles['HeadingTH']))
months = [["มกราคม (ม.ค.)", "กุมภาพันธ์ (ก.พ.)", "มีนาคม (มี.ค.)"],
          ["เมษายน (เม.ย.)", "พฤษภาคม (พ.ค.)", "มิถุนายน (มิ.ย.)"],
          ["กรกฎาคม (ก.ค.)", "สิงหาคม (ส.ค.)", "กันยายน (ก.ย.)"],
          ["ตุลาคม (ต.ค.)", "พฤศจิกายน (พ.ย.)", "ธันวาคม (ธ.ค.)"]]
t = Table(months, colWidths=[160]*3)
t.setStyle(TableStyle([("BOX",(0,0),(-1,-1),1,colors.black),
                       ("INNERGRID",(0,0),(-1,-1),0.5,colors.grey)]))
story.append(t)

docpdf = SimpleDocTemplate(pdf_path, pagesize=A4)
docpdf.build(story)

print("สร้างไฟล์เรียบร้อยแล้วที่:", OUT_DIR)
